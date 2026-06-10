"""
Career-Ops MVP — Enhanced Streamlit App
Tabs: CV Profile | Job Scout | JD Pack | Tracker | Gmail Sync | Settings

Deployment: Render (render.yaml) or local.
Gmail OAuth: web-based flow — works on Render without a local browser.
Scheduled Gmail sync: every 4 days, tracked in data/sync_schedule.json.
"""

import os
import re
import json
import base64
import concurrent.futures
from pathlib import Path
from io import BytesIO
from datetime import datetime, timedelta

import requests
import streamlit as st
import pdfplumber
from bs4 import BeautifulSoup
from openai import OpenAI
from docx import Document

# ── Optional Apify ──────────────────────────────────────────────────
try:
    from apify_client import ApifyClient
    APIFY_AVAILABLE = True
except ImportError:
    APIFY_AVAILABLE = False

# ── Optional Gmail ──────────────────────────────────────────────────
try:
    from google.oauth2.credentials import Credentials
    from google_auth_oauthlib.flow import Flow
    from googleapiclient.discovery import build as google_build
    GMAIL_AVAILABLE = True
except ImportError:
    GMAIL_AVAILABLE = False

# ── Paths ────────────────────────────────────────────────────────────
DATA_DIR         = Path("data")
DATA_DIR.mkdir(exist_ok=True)
CV_PROFILE_PATH  = DATA_DIR / "cv_profile.json"
TRACKER_PATH     = DATA_DIR / "tracker.json"
SETTINGS_PATH    = DATA_DIR / "settings.json"
GMAIL_TOKEN_PATH = DATA_DIR / "gmail_token.json"
SYNC_SCHED_PATH  = DATA_DIR / "sync_schedule.json"
CV_MD_PATH       = Path("cv.md")

GMAIL_SCOPES = ["https://www.googleapis.com/auth/gmail.readonly"]
SYNC_INTERVAL_DAYS = 4

# ── Helpers ──────────────────────────────────────────────────────────

def load_json(path: Path, default):
    if path.exists():
        try:
            return json.loads(path.read_text())
        except Exception:
            return default
    return default

def save_json(path: Path, data):
    path.write_text(json.dumps(data, indent=2, ensure_ascii=False))

def load_settings():
    env_defaults = {
        "openai_key":      os.getenv("OPENAI_API_KEY", ""),
        "apify_key":       os.getenv("APIFY_API_KEY", ""),
        "gmail_client_id":     os.getenv("GMAIL_CLIENT_ID", ""),
        "gmail_client_secret": os.getenv("GMAIL_CLIENT_SECRET", ""),
    }
    saved = load_json(SETTINGS_PATH, {})
    # Saved settings override env vars (so UI changes persist)
    return {**env_defaults, **{k: v for k, v in saved.items() if v}}

def save_settings(s):
    save_json(SETTINGS_PATH, s)

def get_openai_client():
    key = load_settings().get("openai_key", "")
    return OpenAI(api_key=key) if key else None

def get_apify_key():
    return load_settings().get("apify_key", "")

def get_app_base_url() -> str:
    """Returns the base URL for OAuth redirect. Works on Render and locally."""
    url = load_settings().get("app_base_url", "") or os.getenv("APP_BASE_URL", "")
    if url:
        return url.rstrip("/")
    # Auto-detect from Streamlit request headers when running on Render
    try:
        ctx = st.runtime.scriptrunner.get_script_run_ctx()
        if ctx and hasattr(ctx, "request"):
            host = ctx.request.headers.get("host", "")
            if host:
                return f"https://{host}"
    except Exception:
        pass
    return "http://localhost:8501"

def extract_pdf_text(uploaded_file) -> str:
    with pdfplumber.open(uploaded_file) as pdf:
        return "\n".join(page.extract_text() or "" for page in pdf.pages)

def fetch_url_text(url: str) -> str:
    try:
        headers = {"User-Agent": "Mozilla/5.0 (compatible; career-ops/1.0)"}
        r = requests.get(url, headers=headers, timeout=15)
        r.raise_for_status()
        soup = BeautifulSoup(r.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        return soup.get_text(separator="\n", strip=True)[:8000]
    except Exception as e:
        return f"Could not fetch URL: {e}"

def parse_cv_with_llm(cv_text: str, client: OpenAI) -> dict:
    resp = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[
            {"role": "system", "content": (
                "You are a CV parser. Extract these fields and return ONLY a JSON object: "
                "full_name, current_job_title, years_of_experience (integer), "
                "technical_skills (array of strings), professional_summary (2 sentences max), "
                "target_roles (array of strings inferred from CV), "
                "preferred_locations (array of strings if mentioned, else [])."
            )},
            {"role": "user", "content": cv_text},
        ],
        temperature=0.2, max_tokens=1000,
    )
    raw = re.sub(r"^```json|^```|```$", "", resp.choices[0].message.content.strip(), flags=re.MULTILINE).strip()
    return json.loads(raw)

def score_job(profile: dict, jd: str, client: OpenAI) -> int:
    profile_text = (
        f"Title: {profile.get('current_job_title')}\n"
        f"Experience: {profile.get('years_of_experience')} years\n"
        f"Skills: {', '.join(profile.get('technical_skills', []))}\n"
        f"Summary: {profile.get('professional_summary')}"
    )
    resp = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[
            {"role": "system", "content": (
                "Return ONLY a JSON object {\"score\": <integer 0-100>} representing "
                "how well the candidate matches the job. No explanation."
            )},
            {"role": "user", "content": f"Profile:\n{profile_text}\n\nJob:\n{jd}"},
        ],
        temperature=0.1, max_tokens=50,
    )
    raw = re.sub(r"^```json|^```|```$", "", resp.choices[0].message.content.strip(), flags=re.MULTILINE).strip()
    return json.loads(raw).get("score", 0)

def score_jobs_parallel(jobs: list, profile: dict, client: OpenAI) -> list:
    def _score(job):
        try:
            job["score"] = score_job(profile, job.get("description", ""), client)
        except Exception:
            job["score"] = 0
        return job
    with concurrent.futures.ThreadPoolExecutor(max_workers=5) as ex:
        return list(ex.map(_score, jobs))

def scrape_jobs_apify(query: str, location: str, platforms: list, count: int = 10) -> list:
    key = get_apify_key()
    if not key or not APIFY_AVAILABLE:
        return []
    client = ApifyClient(key)
    jobs = []
    if "LinkedIn" in platforms:
        try:
            run = client.actor("curious_coder/linkedin-jobs-scraper").call(
                run_input={"queries": [{"query": query, "location": location}], "maxResults": count})
            for item in client.dataset(run["defaultDatasetId"]).iterate_items():
                jobs.append({"platform": "LinkedIn", "title": item.get("title", ""),
                    "company": item.get("companyName", ""), "location": item.get("location", ""),
                    "description": item.get("description", ""), "url": item.get("jobUrl", ""),
                    "score": 0, "status": "Saved", "saved_at": datetime.now().isoformat()})
        except Exception as e:
            st.warning(f"LinkedIn: {e}")
    if "Indeed" in platforms:
        try:
            run = client.actor("misceres/indeed-scraper").call(
                run_input={"query": query, "location": location, "maxItems": count})
            for item in client.dataset(run["defaultDatasetId"]).iterate_items():
                jobs.append({"platform": "Indeed", "title": item.get("positionName", ""),
                    "company": item.get("company", ""), "location": item.get("location", ""),
                    "description": item.get("description", ""), "url": item.get("url", ""),
                    "score": 0, "status": "Saved", "saved_at": datetime.now().isoformat()})
        except Exception as e:
            st.warning(f"Indeed: {e}")
    if "Seek" in platforms:
        try:
            run = client.actor("bebity/seek-jobs-scraper").call(
                run_input={"keyword": query, "location": location, "maxItems": count})
            for item in client.dataset(run["defaultDatasetId"]).iterate_items():
                jobs.append({"platform": "Seek", "title": item.get("title", ""),
                    "company": item.get("advertiser", {}).get("description", ""),
                    "location": item.get("location", ""),
                    "description": item.get("teaser", ""), "url": item.get("jobUrl", ""),
                    "score": 0, "status": "Saved", "saved_at": datetime.now().isoformat()})
        except Exception as e:
            st.warning(f"Seek: {e}")
    return jobs

def build_cv_tailor_prompt(cv_text: str, jd: str, company: str) -> tuple:
    system = (
        "You are an expert job application strategist and ATS optimization specialist. "
        "Use the base CV as canonical truth. NEVER invent claims, metrics, tools, dates, or responsibilities. "
        "Return markdown with exactly these sections:\n"
        "## Fitment Score\n## ATS-Tailored CV\n## Cover Letter\n## Interview Prep (3 likely questions + answers)"
    )
    user = f"Base CV:\n{cv_text}\n\nJob Description:\n{jd}\n\nCompany Context:\n{company or 'Not provided'}"
    return system, user

def markdown_to_docx(md: str) -> bytes:
    doc = Document()
    for line in md.splitlines():
        if line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif re.match(r"^[-*]\s+", line):
            doc.add_paragraph(re.sub(r"^[-*]\s+", "", line).strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s+", line):
            doc.add_paragraph(re.sub(r"^\d+\.\s+", "", line).strip(), style="List Number")
        elif line.strip():
            doc.add_paragraph(line)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

def score_badge(score: int) -> str:
    if score >= 70: return f"🟢 {score}%"
    elif score >= 40: return f"🟡 {score}%"
    return f"🔴 {score}%"

# ── Gmail helpers ────────────────────────────────────────────────────

REJECTION_KEYWORDS = ["unfortunately", "not moving forward", "not selected", "other candidates",
    "decided to move on", "not a fit", "position has been filled", "we won't be moving",
    "not progressing", "no longer considering", "regret to inform", "not successful", "we have decided"]
INTERVIEW_KEYWORDS = ["interview", "schedule a call", "next steps", "we'd like to speak",
    "shortlisted", "assessment", "technical screen", "hiring manager"]
OFFER_KEYWORDS     = ["offer letter", "pleased to offer", "congratulations", "we are delighted",
    "formal offer", "compensation package"]

def classify_email(subject: str, body: str) -> str | None:
    text = (subject + " " + body).lower()
    if any(k in text for k in OFFER_KEYWORDS):     return "Offer"
    if any(k in text for k in INTERVIEW_KEYWORDS): return "Interview"
    if any(k in text for k in REJECTION_KEYWORDS): return "Rejected"
    return None

def get_gmail_service():
    token_data = load_json(GMAIL_TOKEN_PATH, None)
    if not token_data:
        return None
    try:
        creds = Credentials.from_authorized_user_info(token_data, GMAIL_SCOPES)
        return google_build("gmail", "v1", credentials=creds)
    except Exception:
        return None

def fetch_job_emails(service, max_results: int = 50) -> list:
    results = []
    try:
        query = "subject:(application OR interview OR offer OR opportunity OR position OR role OR hiring)"
        msgs = service.users().messages().list(userId="me", q=query, maxResults=max_results).execute().get("messages", [])
        for msg in msgs:
            m = service.users().messages().get(userId="me", id=msg["id"], format="full").execute()
            headers = {h["name"]: h["value"] for h in m["payload"].get("headers", [])}
            subject = headers.get("Subject", "")
            sender  = headers.get("From", "")
            date    = headers.get("Date", "")
            body = ""
            parts = m["payload"].get("parts", [])
            if parts:
                for part in parts:
                    if part.get("mimeType") == "text/plain":
                        data = part.get("body", {}).get("data", "")
                        if data:
                            body = base64.urlsafe_b64decode(data).decode("utf-8", errors="ignore")
                            break
            else:
                data = m["payload"].get("body", {}).get("data", "")
                if data:
                    body = base64.urlsafe_b64decode(data).decode("utf-8", errors="ignore")
            classification = classify_email(subject, body)
            if classification:
                results.append({"id": msg["id"], "subject": subject, "sender": sender,
                    "date": date, "classification": classification, "body_preview": body[:300]})
    except Exception as e:
        st.error(f"Gmail fetch error: {e}")
    return results

def match_email_to_tracker(email: dict, tracker: list) -> int | None:
    sender_lower  = email["sender"].lower()
    subject_lower = email["subject"].lower()
    for i, job in enumerate(tracker):
        company = job.get("company", "").lower()
        if company and len(company) > 2:
            if company in sender_lower or company in subject_lower:
                return i
    return None

def run_gmail_sync(service) -> dict:
    """Run a full Gmail sync and return a summary dict."""
    tracker_data = load_json(TRACKER_PATH, [])
    emails = fetch_job_emails(service, 100)
    STATUSES = ["Saved", "Applied", "Screening", "Interview", "Offer", "Rejected"]
    status_rank = {s: i for i, s in enumerate(STATUSES)}
    updates = 0
    for email in emails:
        idx = match_email_to_tracker(email, tracker_data)
        if idx is not None:
            old_status = tracker_data[idx].get("status", "Saved")
            new_status = email["classification"]
            if status_rank.get(new_status, 0) > status_rank.get(old_status, 0):
                tracker_data[idx]["status"] = new_status
                tracker_data[idx]["updated_at"] = datetime.now().isoformat()
                tracker_data[idx]["email_match"] = email["subject"]
                updates += 1
    if updates:
        save_json(TRACKER_PATH, tracker_data)
    # Record sync time
    save_json(SYNC_SCHED_PATH, {"last_sync": datetime.now().isoformat(), "interval_days": SYNC_INTERVAL_DAYS})
    return {"emails_found": len(emails), "tracker_updates": updates}

def check_and_auto_sync():
    """Auto-trigger Gmail sync if 4 days have passed since last sync."""
    if not GMAIL_AVAILABLE:
        return
    sched = load_json(SYNC_SCHED_PATH, {})
    last_sync_str = sched.get("last_sync")
    if last_sync_str:
        last_sync = datetime.fromisoformat(last_sync_str)
        if datetime.now() - last_sync < timedelta(days=SYNC_INTERVAL_DAYS):
            return  # Not due yet
    service = get_gmail_service()
    if not service:
        return
    result = run_gmail_sync(service)
    if result["tracker_updates"] > 0:
        st.toast(f"📧 Gmail auto-sync: {result['tracker_updates']} tracker updates", icon="✅")

# ── Run scheduled sync on every app load ────────────────────────────
check_and_auto_sync()

# ── App ──────────────────────────────────────────────────────────────

st.set_page_config(page_title="Career-Ops", page_icon="🧭", layout="wide")
st.markdown("""
<style>
.block-container { padding-top: 1.5rem; }
.stTabs [data-baseweb="tab"] { font-size: 15px; font-weight: 600; }
</style>
""", unsafe_allow_html=True)

st.title("🧭 Career-Ops")
st.caption("AI-powered job scout, CV tailor, and application tracker — personal MVP")

# Show next scheduled sync time
sched = load_json(SYNC_SCHED_PATH, {})
if sched.get("last_sync"):
    next_sync = datetime.fromisoformat(sched["last_sync"]) + timedelta(days=SYNC_INTERVAL_DAYS)
    st.caption(f"📧 Next Gmail auto-sync: **{next_sync.strftime('%d %b %Y')}**")

tabs = st.tabs(["👤 CV Profile", "🔍 Job Scout", "✍️ JD Pack", "📋 Tracker", "📧 Gmail Sync", "⚙️ Settings"])

# ─────────────────────────────────────────────────────────────────────
# TAB 1 — CV PROFILE
# ─────────────────────────────────────────────────────────────────────
with tabs[0]:
    st.header("CV Profile")
    st.caption("Upload your CV once. All agents use this as the source of truth.")
    profile = load_json(CV_PROFILE_PATH, {})
    uploaded = st.file_uploader("Upload your CV (PDF)", type=["pdf"])
    if uploaded:
        client = get_openai_client()
        if not client:
            st.error("Set your OpenAI API key in ⚙️ Settings first.")
        else:
            with st.spinner("Parsing CV..."):
                cv_text = extract_pdf_text(uploaded)
                try:
                    profile = parse_cv_with_llm(cv_text, client)
                    profile["raw_text"] = cv_text
                    save_json(CV_PROFILE_PATH, profile)
                    CV_MD_PATH.write_text(cv_text)
                    st.success("CV parsed and saved.")
                except Exception as e:
                    st.error(f"Parsing failed: {e}")
    if profile:
        col1, col2 = st.columns(2)
        with col1:
            st.subheader(profile.get("full_name", "—"))
            st.write(f"**Title:** {profile.get('current_job_title', '—')}")
            st.write(f"**Experience:** {profile.get('years_of_experience', '—')} years")
            st.write(f"**Target Roles:** {', '.join(profile.get('target_roles', []))}")
            st.write(f"**Preferred Locations:** {', '.join(profile.get('preferred_locations', []))}")
        with col2:
            st.write("**Technical Skills:**")
            st.write(", ".join(profile.get("technical_skills", [])) or "—")
            st.write("**Summary:**")
            st.info(profile.get("professional_summary", "—"))
    else:
        st.info("No CV profile yet. Upload your CV above.")

# ─────────────────────────────────────────────────────────────────────
# TAB 2 — JOB SCOUT
# ─────────────────────────────────────────────────────────────────────
with tabs[1]:
    st.header("Job Scout")
    profile = load_json(CV_PROFILE_PATH, {})
    if not profile:
        st.warning("Upload your CV in the **CV Profile** tab first.")
    else:
        scout_mode = st.radio(
            "How do you want to find jobs?",
            ["🤖 Auto-fetch via Apify (LinkedIn/Indeed/Seek)", "📋 Paste JD text manually", "🔗 Paste job URL"],
            horizontal=True,
        )

        if scout_mode == "🤖 Auto-fetch via Apify (LinkedIn/Indeed/Seek)":
            if not get_apify_key():
                st.warning("No Apify key set. Go to ⚙️ Settings to add one, or use the manual options.")
            else:
                col1, col2, col3 = st.columns([2, 2, 1])
                with col1:
                    default_role = profile.get("target_roles", [""])[0] if profile.get("target_roles") else ""
                    query = st.text_input("Job Title / Keywords", value=default_role)
                with col2:
                    default_loc = profile.get("preferred_locations", [""])[0] if profile.get("preferred_locations") else ""
                    location = st.text_input("Location", value=default_loc)
                with col3:
                    count = st.number_input("Results per platform", min_value=3, max_value=20, value=8)
                platforms = st.multiselect("Platforms", ["LinkedIn", "Indeed", "Seek"], default=["LinkedIn", "Indeed"])
                if st.button("🔍 Fetch & Score Jobs", type="primary"):
                    client = get_openai_client()
                    if not client:
                        st.error("Set your OpenAI API key in ⚙️ Settings.")
                    elif not query or not location:
                        st.error("Enter a job title and location.")
                    else:
                        with st.spinner("Scraping job boards..."):
                            jobs = scrape_jobs_apify(query, location, platforms, int(count))
                        if not jobs:
                            st.error("No jobs returned. Check your Apify key and actor availability.")
                        else:
                            with st.spinner(f"Scoring {len(jobs)} jobs..."):
                                jobs = score_jobs_parallel(jobs, profile, client)
                            jobs.sort(key=lambda j: j["score"], reverse=True)
                            st.session_state["scout_jobs"] = jobs
                            st.success(f"Found and scored {len(jobs)} jobs.")

        elif scout_mode == "📋 Paste JD text manually":
            manual_title    = st.text_input("Job Title")
            manual_company  = st.text_input("Company Name")
            manual_location = st.text_input("Location")
            manual_url      = st.text_input("Job URL (optional)")
            manual_jd       = st.text_area("Paste Job Description", height=250)
            if st.button("⚡ Score & Add", type="primary"):
                client = get_openai_client()
                if not client:
                    st.error("Set your OpenAI API key in ⚙️ Settings.")
                elif not manual_jd.strip():
                    st.error("Paste a job description first.")
                else:
                    with st.spinner("Scoring..."):
                        score = score_job(profile, manual_jd, client)
                    job = {"platform": "Manual", "title": manual_title, "company": manual_company,
                        "location": manual_location, "description": manual_jd,
                        "url": manual_url or "#", "score": score,
                        "status": "Saved", "saved_at": datetime.now().isoformat()}
                    existing = st.session_state.get("scout_jobs", [])
                    existing.insert(0, job)
                    st.session_state["scout_jobs"] = existing
                    st.success(f"Added — {score_badge(score)}")

        elif scout_mode == "🔗 Paste job URL":
            url_input      = st.text_input("Job URL (LinkedIn, Indeed, Seek, or any job board)")
            manual_title   = st.text_input("Job Title (optional override)")
            manual_company = st.text_input("Company Name (optional override)")
            if st.button("🌐 Fetch, Score & Add", type="primary"):
                client = get_openai_client()
                if not client:
                    st.error("Set your OpenAI API key in ⚙️ Settings.")
                elif not url_input.strip():
                    st.error("Enter a URL.")
                else:
                    with st.spinner("Fetching job page..."):
                        jd_text = fetch_url_text(url_input)
                    if jd_text.startswith("Could not fetch"):
                        st.error(jd_text + " — try pasting the JD text manually instead.")
                    else:
                        with st.spinner("Scoring..."):
                            score = score_job(profile, jd_text, client)
                        job = {"platform": "URL", "title": manual_title or "From URL",
                            "company": manual_company or "Unknown", "location": "",
                            "description": jd_text[:3000], "url": url_input, "score": score,
                            "status": "Saved", "saved_at": datetime.now().isoformat()}
                        existing = st.session_state.get("scout_jobs", [])
                        existing.insert(0, job)
                        st.session_state["scout_jobs"] = existing
                        st.success(f"Fetched and scored {score_badge(score)}")

        scout_jobs = st.session_state.get("scout_jobs", [])
        if scout_jobs:
            st.divider()
            st.subheader(f"Results ({len(scout_jobs)} jobs)")
            tracker_data = load_json(TRACKER_PATH, [])
            existing_urls = {j["url"] for j in tracker_data}
            for job in scout_jobs:
                with st.expander(f"{score_badge(job['score'])}  **{job['title']}** — {job['company']}  [{job['platform']}]"):
                    st.write(f"📍 {job['location']}  |  🔗 [{job['url']}]({job['url']})")
                    desc = job.get("description", "")
                    st.write((desc[:500] + "...") if len(desc) > 500 else desc)
                    c1, c2 = st.columns(2)
                    with c1:
                        if st.button("💾 Save to Tracker", key=f"save_{job['url']}"):
                            if job["url"] not in existing_urls:
                                tracker_data.append(job)
                                save_json(TRACKER_PATH, tracker_data)
                                existing_urls.add(job["url"])
                                st.success("Saved.")
                            else:
                                st.info("Already in Tracker.")
                    with c2:
                        if st.button("✍️ Tailor CV", key=f"tailor_{job['url']}"):
                            st.session_state["jdpack_jd"]      = job.get("description", "")
                            st.session_state["jdpack_company"] = job.get("company", "")
                            st.info("JD loaded — switch to ✍️ JD Pack tab.")

# ─────────────────────────────────────────────────────────────────────
# TAB 3 — JD PACK
# ─────────────────────────────────────────────────────────────────────
with tabs[2]:
    st.header("JD Pack — CV Tailor & Cover Letter")
    client = get_openai_client()
    if not client:
        st.error("Set your OpenAI API key in ⚙️ Settings.")
    else:
        jd_text = st.text_area("Job Description", height=220,
            value=st.session_state.get("jdpack_jd", ""),
            placeholder="Paste JD here or fetch from Job Scout tab...")
        company_profile = st.text_area("Company Context (optional)", height=100,
            value=st.session_state.get("jdpack_company", ""))
        col1, col2 = st.columns(2)
        with col1:
            model = st.selectbox("Model", ["gpt-4.1-mini", "gpt-4.1", "gpt-4o"])
        with col2:
            max_tokens = st.slider("Max output tokens", 1000, 4000, 2000, 200)
        if st.button("⚡ Generate Application Pack", type="primary"):
            if not jd_text.strip():
                st.error("Paste a job description first.")
            else:
                cv_text = ""
                if CV_MD_PATH.exists():
                    cv_text = CV_MD_PATH.read_text()
                elif CV_PROFILE_PATH.exists():
                    cv_text = load_json(CV_PROFILE_PATH, {}).get("raw_text", "")
                if not cv_text:
                    st.error("No CV found. Upload your CV in the CV Profile tab first.")
                else:
                    system_p, user_p = build_cv_tailor_prompt(cv_text, jd_text, company_profile)
                    with st.spinner("Generating..."):
                        resp = client.chat.completions.create(
                            model=model,
                            messages=[{"role": "system", "content": system_p}, {"role": "user", "content": user_p}],
                            max_tokens=max_tokens, temperature=0.3)
                        output = resp.choices[0].message.content.strip()
                    st.session_state["jdpack_output"] = output
                    st.success("Generated.")
        output = st.session_state.get("jdpack_output", "")
        if output:
            st.markdown(output)
            d1, d2 = st.columns(2)
            with d1:
                st.download_button("⬇️ Markdown", data=output.encode(), file_name="application-pack.md",
                    mime="text/markdown", use_container_width=True)
            with d2:
                st.download_button("⬇️ Word (.docx)", data=markdown_to_docx(output),
                    file_name="application-pack.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True)

# ─────────────────────────────────────────────────────────────────────
# TAB 4 — TRACKER
# ─────────────────────────────────────────────────────────────────────
with tabs[3]:
    st.header("Application Tracker")
    STATUSES = ["Saved", "Applied", "Screening", "Interview", "Offer", "Rejected"]
    STATUS_ICONS = {"Saved": "💾", "Applied": "📤", "Screening": "🔍",
                    "Interview": "🗣️", "Offer": "🎉", "Rejected": "❌"}
    tracker_data = load_json(TRACKER_PATH, [])
    if not tracker_data:
        st.info("No jobs tracked yet. Save jobs from the Job Scout tab.")
    else:
        cols = st.columns(len(STATUSES))
        for i, s in enumerate(STATUSES):
            cols[i].metric(f"{STATUS_ICONS[s]} {s}", sum(1 for j in tracker_data if j.get("status") == s))
        st.divider()
        changed = False
        for status in STATUSES:
            jobs_in_status = [j for j in tracker_data if j.get("status") == status]
            if not jobs_in_status:
                continue
            st.subheader(f"{STATUS_ICONS[status]} {status} ({len(jobs_in_status)})")
            for idx, job in enumerate(tracker_data):
                if job.get("status") != status:
                    continue
                with st.expander(f"**{job.get('title','—')}** — {job.get('company','—')}  {score_badge(job.get('score',0))}"):
                    st.write(f"📍 {job.get('location','—')}  |  🌐 {job.get('platform','—')}")
                    st.write(f"🔗 [{job.get('url','')}]({job.get('url','')})")
                    if job.get("email_match"):
                        st.info(f"📧 Auto-updated from Gmail: *{job['email_match']}*")
                    new_status = st.selectbox("Update Status", STATUSES,
                        index=STATUSES.index(job.get("status", "Saved")),
                        key=f"status_{idx}_{job.get('url', idx)}")
                    if new_status != job.get("status"):
                        tracker_data[idx]["status"] = new_status
                        tracker_data[idx]["updated_at"] = datetime.now().isoformat()
                        changed = True
                    if st.button("🗑️ Remove", key=f"del_{idx}_{job.get('url', idx)}"):
                        tracker_data.pop(idx)
                        save_json(TRACKER_PATH, tracker_data)
                        st.rerun()
        if changed:
            save_json(TRACKER_PATH, tracker_data)
            st.success("Tracker updated.")
        st.divider()
        st.subheader("📊 Quick Stats")
        total   = len(tracker_data)
        applied = sum(1 for j in tracker_data if j.get("status") not in ["Saved", "Rejected"])
        avg_score = round(sum(j.get("score", 0) for j in tracker_data) / total, 1) if total else 0
        responded = sum(1 for j in tracker_data if j.get("status") in ["Screening", "Interview", "Offer"])
        response_rate = round(responded / applied * 100, 1) if applied else 0
        m1, m2, m3, m4 = st.columns(4)
        m1.metric("Total Tracked", total)
        m2.metric("Applied", applied)
        m3.metric("Response Rate", f"{response_rate}%")
        m4.metric("Avg Match Score", f"{avg_score}%")

# ─────────────────────────────────────────────────────────────────────
# TAB 5 — GMAIL SYNC
# ─────────────────────────────────────────────────────────────────────
with tabs[4]:
    st.header("📧 Gmail Sync")
    st.caption("Auto-syncs every 4 days. Detects rejection, interview, and offer emails and updates your tracker.")

    if not GMAIL_AVAILABLE:
        st.error("Gmail libraries not installed. Run: `pip install google-auth google-auth-oauthlib google-api-python-client`")
    else:
        settings = load_settings()
        client_id     = settings.get("gmail_client_id", "")
        client_secret = settings.get("gmail_client_secret", "")

        if not client_id or not client_secret:
            st.warning("Add your Gmail OAuth credentials in ⚙️ Settings first.")
            with st.expander("📖 How to get Gmail OAuth credentials (free, 5 min)"):
                st.markdown("""
1. Go to [Google Cloud Console](https://console.cloud.google.com/) → **New Project**
2. Search for **Gmail API** → Enable it
3. Go to **APIs & Services → Credentials → Create Credentials → OAuth 2.0 Client ID**
4. Application type: **Web application**
5. Under **Authorised redirect URIs**, add your Render URL + `/` e.g.:
   `https://career-ops-streamlit.onrender.com/`
6. Copy the **Client ID** and **Client Secret** → paste into ⚙️ Settings
7. Also set **APP_BASE_URL** in Render environment variables to your app URL
                """)
        else:
            token_data    = load_json(GMAIL_TOKEN_PATH, None)
            gmail_service = get_gmail_service() if token_data else None

            # ── CONNECT FLOW ──
            if not gmail_service:
                st.info("Gmail is not connected yet.")

                # Handle OAuth callback — Streamlit passes ?code= in query params
                query_params = st.query_params
                auth_code = query_params.get("code")

                if auth_code:
                    # Exchange code for token
                    try:
                        base_url = get_app_base_url()
                        flow = Flow.from_client_config(
                            {"web": {
                                "client_id": client_id, "client_secret": client_secret,
                                "auth_uri": "https://accounts.google.com/o/oauth2/auth",
                                "token_uri": "https://oauth2.googleapis.com/token",
                                "redirect_uris": [base_url + "/"],
                            }},
                            scopes=GMAIL_SCOPES,
                            redirect_uri=base_url + "/",
                        )
                        flow.fetch_token(code=auth_code)
                        creds = flow.credentials
                        save_json(GMAIL_TOKEN_PATH, json.loads(creds.to_json()))
                        st.query_params.clear()
                        st.success("✅ Gmail connected successfully!")
                        st.rerun()
                    except Exception as e:
                        st.error(f"OAuth error: {e}. Try connecting again.")

                else:
                    # Show connect button
                    if st.button("🔐 Connect Gmail", type="primary"):
                        try:
                            base_url = get_app_base_url()
                            flow = Flow.from_client_config(
                                {"web": {
                                    "client_id": client_id, "client_secret": client_secret,
                                    "auth_uri": "https://accounts.google.com/o/oauth2/auth",
                                    "token_uri": "https://oauth2.googleapis.com/token",
                                    "redirect_uris": [base_url + "/"],
                                }},
                                scopes=GMAIL_SCOPES,
                                redirect_uri=base_url + "/",
                            )
                            auth_url, _ = flow.authorization_url(prompt="consent", access_type="offline")
                            st.markdown(f"""
### Click to authorise Gmail access

**[👉 Authorise Gmail →]({auth_url})**

You will be redirected back to this app automatically after authorising.
                            """)
                        except Exception as e:
                            st.error(f"Could not build auth URL: {e}")

            # ── CONNECTED ──
            else:
                st.success("✅ Gmail connected.")
                sched = load_json(SYNC_SCHED_PATH, {})
                if sched.get("last_sync"):
                    last = datetime.fromisoformat(sched["last_sync"])
                    nxt  = last + timedelta(days=SYNC_INTERVAL_DAYS)
                    st.write(f"**Last sync:** {last.strftime('%d %b %Y %H:%M')}  |  **Next auto-sync:** {nxt.strftime('%d %b %Y')}")

                col1, col2 = st.columns(2)
                with col1:
                    max_emails = st.number_input("Emails to scan", min_value=10, max_value=200, value=50)
                with col2:
                    st.write("")
                    st.write("")
                    run_sync = st.button("🔄 Sync Now", type="primary")

                if run_sync:
                    with st.spinner("Scanning Gmail..."):
                        result = run_gmail_sync(gmail_service)
                    st.success(f"✅ Found {result['emails_found']} job emails — {result['tracker_updates']} tracker updates.")

                # Show recent detected emails
                if st.button("👁️ Preview recent job emails"):
                    with st.spinner("Fetching..."):
                        emails = fetch_job_emails(gmail_service, int(max_emails))
                    tracker_data = load_json(TRACKER_PATH, [])
                    for email in emails:
                        icon = {"Rejected": "❌", "Interview": "🗣️", "Offer": "🎉"}.get(email["classification"], "📧")
                        with st.expander(f"{icon} {email['classification']} — {email['subject'][:60]}"):
                            st.write(f"**From:** {email['sender']}  |  **Date:** {email['date']}")
                            st.write(email["body_preview"])
                            idx = match_email_to_tracker(email, tracker_data)
                            if idx is not None:
                                st.success(f"Matched: **{tracker_data[idx].get('company')} — {tracker_data[idx].get('title')}**")
                            else:
                                st.warning("No tracker match found.")

                if st.button("🔌 Disconnect Gmail"):
                    GMAIL_TOKEN_PATH.unlink(missing_ok=True)
                    st.success("Disconnected.")
                    st.rerun()

# ─────────────────────────────────────────────────────────────────────
# TAB 6 — SETTINGS
# ─────────────────────────────────────────────────────────────────────
with tabs[5]:
    st.header("⚙️ Settings")
    settings = load_settings()

    st.subheader("OpenAI")
    openai_key = st.text_input("OpenAI API Key", value=settings.get("openai_key", ""), type="password")

    st.subheader("Apify (optional)")
    apify_key = st.text_input("Apify API Key", value=settings.get("apify_key", ""), type="password",
        help="Leave blank to use manual JD input only.")

    st.subheader("Gmail OAuth (optional)")
    gmail_client_id     = st.text_input("Gmail OAuth Client ID",     value=settings.get("gmail_client_id", ""),     type="password")
    gmail_client_secret = st.text_input("Gmail OAuth Client Secret", value=settings.get("gmail_client_secret", ""), type="password")
    app_base_url        = st.text_input("App Base URL (your Render URL)",
        value=settings.get("app_base_url", os.getenv("APP_BASE_URL", "")),
        placeholder="https://career-ops-streamlit.onrender.com",
        help="Required for Gmail OAuth redirect. Set this to your Render app URL.")

    if st.button("💾 Save Settings", type="primary"):
        save_settings({"openai_key": openai_key, "apify_key": apify_key,
            "gmail_client_id": gmail_client_id, "gmail_client_secret": gmail_client_secret,
            "app_base_url": app_base_url})
        st.success("Settings saved.")

    st.divider()
    st.subheader("Data Management")
    c1, c2, c3 = st.columns(3)
    with c1:
        if st.button("🗑️ Clear CV Profile"):
            CV_PROFILE_PATH.unlink(missing_ok=True)
            st.success("CV profile cleared.")
    with c2:
        if st.button("🗑️ Clear Tracker"):
            TRACKER_PATH.unlink(missing_ok=True)
            st.success("Tracker cleared.")
    with c3:
        if st.button("🗑️ Disconnect Gmail"):
            GMAIL_TOKEN_PATH.unlink(missing_ok=True)
            st.success("Gmail disconnected.")
