import os
import re
import json
from pathlib import Path
from io import BytesIO

import streamlit as st
from openai import OpenAI
from docx import Document


MEMORY_PATH = Path("data/learning-memory.json")

DEFAULT_REASONING_PROFILE = [
    "Reason deeply about strategic fit before rewriting content.",
    "Identify role archetype, required operating context, and likely interview risk factors.",
    "Explicitly map candidate strengths and gaps against JD demands.",
    "When asked to improve score (e.g., 4.2 to 4.8), prioritize repositioning and stronger evidence selection before adding wording changes.",
    "For industrial or operations-heavy roles, bias language toward operational leadership, governance, workforce planning, multi-site partnership, and leadership maturity.",
    "Reduce AI/startup-heavy vocabulary when role context requires operational credibility.",
    "Keep output concise, but keep internal reasoning rigorous.",
]


def tokenize_words(text: str) -> set[str]:
    return set(re.findall(r"[a-zA-Z][a-zA-Z0-9+.#/-]{1,}", text.lower()))


def relevant_cv_slice(cv_text: str, jd_text: str, company_profile: str, max_lines: int = 140) -> str:
    lines = [ln.rstrip() for ln in cv_text.splitlines() if ln.strip()]
    context_terms = tokenize_words(jd_text + "\n" + company_profile)
    if not context_terms:
        return "\n".join(lines[:max_lines])

    scored = []
    for idx, line in enumerate(lines):
        words = tokenize_words(line)
        overlap = len(words.intersection(context_terms))
        bonus = 2 if line.startswith("#") or line.endswith(":") else 0
        scored.append((overlap + bonus, idx, line))

    top = [row for row in sorted(scored, key=lambda x: (x[0], -x[1]), reverse=True) if row[0] > 0]
    keep = set()
    for _, idx, _ in top[:max_lines]:
        keep.add(idx)
        if idx > 0:
            keep.add(idx - 1)
        if idx + 1 < len(lines):
            keep.add(idx + 1)

    selected = [lines[i] for i in sorted(keep)]
    if not selected:
        selected = lines[:max_lines]
    return "\n".join(selected[: max_lines * 2])


def load_learning_memory() -> dict:
    if not MEMORY_PATH.exists():
        return {"notes": [], "reasoning_profile": DEFAULT_REASONING_PROFILE}
    try:
        data = json.loads(MEMORY_PATH.read_text(encoding="utf-8"))
        notes = data.get("notes", [])
        profile = data.get("reasoning_profile", DEFAULT_REASONING_PROFILE)
        if not isinstance(notes, list):
            notes = []
        if not isinstance(profile, list) or not profile:
            profile = DEFAULT_REASONING_PROFILE
        return {"notes": notes, "reasoning_profile": profile}
    except Exception:
        return {"notes": [], "reasoning_profile": DEFAULT_REASONING_PROFILE}


def save_learning_memory(notes: list[str], reasoning_profile: list[str]) -> None:
    MEMORY_PATH.parent.mkdir(parents=True, exist_ok=True)
    payload = {"notes": notes[-200:], "reasoning_profile": reasoning_profile}
    MEMORY_PATH.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")


def build_prompts(
    cv_text: str,
    jd_text: str,
    company_profile: str,
    rewrite_request: str,
    learning_notes: list[str],
    reasoning_profile: list[str],
    compact_mode: bool,
) -> tuple[str, str]:
    output_contract = [
        "Return markdown only with exactly these sections in order:",
        "## 1) Comprehensive Scoring",
        "## 2) Customized CV (Copy-Paste)",
        "## 3) Customized Cover Letter (Copy-Paste)",
    ]

    style_rules = [
        "Use cv.md context as canonical truth. Never invent claims, metrics, tools, dates, or responsibilities.",
        "For scoring, provide A-G style summary and final score out of 5 with risks and mitigations.",
        "For CV, produce ATS-friendly markdown with role-aligned summary, skills, and tailored experience bullets.",
        "For cover letter, keep concise, specific to role/company, and practical for direct copy-paste.",
        "If user asks to improve score (e.g., 4.2 to 4.8), optimize framing, alignment, and prioritization without fabricating facts.",
        "CV completeness rule: include ALL experience sections present in the base CV; never omit a role.",
        "CV section order rule: header/contact, EXECUTIVE SUMMARY, KEY IMPACT, PROFESSIONAL EXPERIENCE, CORE SKILLS, EDUCATION, CERTIFICATIONS.",
        "Place CORE SKILLS only after PROFESSIONAL EXPERIENCE.",
        "If KEY IMPACT is not explicit in base CV, derive it only from existing evidence in base CV without inventing facts.",
    ]
    if compact_mode:
        style_rules.append("Be concise and avoid unnecessary verbosity to reduce token usage.")

    learn_block = "\n".join(f"- {note}" for note in learning_notes[-12:]) if learning_notes else "- none yet"
    reasoning_block = "\n".join(f"- {item}" for item in reasoning_profile) if reasoning_profile else "- none yet"

    system_prompt = "\n".join(
        [
            "You are an expert job application strategist.",
            *output_contract,
            *style_rules,
            "Reasoning behavior to follow internally (do not expose chain-of-thought):",
            reasoning_block,
            "Apply this persistent user preference memory:",
            learn_block,
        ]
    )

    user_prompt = "\n".join(
        [
            "Base CV context:",
            cv_text,
            "",
            "Job Description:",
            jd_text,
            "",
            "Company Profile:",
            company_profile if company_profile.strip() else "(not provided)",
            "",
            "Rewrite objective / feedback from user:",
            rewrite_request if rewrite_request.strip() else "(none)",
        ]
    )
    return system_prompt, user_prompt


def markdown_to_docx_bytes(markdown_text: str) -> bytes:
    doc = Document()
    for raw in markdown_text.splitlines():
        line = raw.rstrip()
        if not line.strip():
            doc.add_paragraph("")
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith("- ") or line.startswith("• "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
            continue
        if re.match(r"^\d+\.\s+", line):
            text = re.sub(r"^\d+\.\s+", "", line).strip()
            doc.add_paragraph(text, style="List Number")
            continue
        doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


st.set_page_config(page_title="Career-Ops JD Pack", page_icon="🧭", layout="wide")
st.title("Career-Ops: JD -> Scoring + CV + Cover Letter")
st.caption("Paste JD + company profile, generate, refine, and copy-paste.")

cv_path = Path("cv.md")
if not cv_path.exists():
    st.error("`cv.md` not found in repo root. Add your canonical base CV first.")
    st.stop()

if "learning_notes" not in st.session_state:
    memory = load_learning_memory()
    st.session_state.learning_notes = memory["notes"]
    st.session_state.reasoning_profile = memory["reasoning_profile"]
if "latest_output" not in st.session_state:
    st.session_state.latest_output = ""

api_key = os.getenv("OPENAI_API_KEY", "").strip()
if not api_key:
    st.warning("Set `OPENAI_API_KEY` in environment variables (Render -> Environment).")

jd_text = st.text_area("Job Description (required)", height=240, placeholder="Paste JD here...")
company_profile = st.text_area(
    "Company Profile from LinkedIn (optional)",
    height=150,
    placeholder="Paste company profile/context here...",
)
rewrite_request = st.text_area(
    "Rewrite Prompt / Feedback (optional)",
    height=120,
    placeholder="Example: Raise fit from 4.2 to 4.8 by emphasizing HR shared services transformation and compliance leadership.",
)

col1, col2, col3 = st.columns(3)
with col1:
    model = st.text_input("Model", value=os.getenv("OPENAI_MODEL", "gpt-4.1"))
with col2:
    compact_mode = st.checkbox("Token-optimized output", value=True)
with col3:
    use_relevant_slice = st.checkbox("Use relevant CV slice", value=False)

max_output_tokens = st.slider("Max output tokens", min_value=800, max_value=5000, value=1800, step=100)

btn_col1, btn_col2, btn_col3 = st.columns(3)
generate = btn_col1.button("Generate", type="primary", use_container_width=True)
regenerate = btn_col2.button("Regenerate with Feedback", use_container_width=True)
clear_learning = btn_col3.button("Clear Learning Memory", use_container_width=True)

if clear_learning:
    st.session_state.learning_notes = []
    save_learning_memory(st.session_state.learning_notes, st.session_state.reasoning_profile)
    st.success("Learning memory cleared for this session.")

should_run = generate or regenerate
if should_run:
    if not jd_text.strip():
        st.error("Please paste the job description.")
        st.stop()
    if not api_key:
        st.error("Missing `OPENAI_API_KEY` environment variable.")
        st.stop()

    cv_full = cv_path.read_text(encoding="utf-8")
    cv_context = relevant_cv_slice(cv_full, jd_text, company_profile) if use_relevant_slice else cv_full

    if rewrite_request.strip():
        st.session_state.learning_notes.append(rewrite_request.strip())
        save_learning_memory(st.session_state.learning_notes, st.session_state.reasoning_profile)

    system_prompt, user_prompt = build_prompts(
        cv_context,
        jd_text,
        company_profile,
        rewrite_request,
        st.session_state.learning_notes,
        st.session_state.reasoning_profile,
        compact_mode,
    )

    client = OpenAI(api_key=api_key)
    with st.spinner("Generating..."):
        response = client.responses.create(
            model=model,
            max_output_tokens=max_output_tokens,
            input=[
                {"role": "system", "content": [{"type": "input_text", "text": system_prompt}]},
                {"role": "user", "content": [{"type": "input_text", "text": user_prompt}]},
            ],
        )
        output = (response.output_text or "").strip()

    if not output:
        st.error("Model returned empty output. Try again.")
        st.stop()

    st.session_state.latest_output = output
    st.success("Generated.")

if st.session_state.latest_output:
    st.markdown(st.session_state.latest_output)
    d1, d2 = st.columns(2)
    with d1:
        st.download_button(
            label="Download as Markdown",
            data=st.session_state.latest_output.encode("utf-8"),
            file_name="application-pack.md",
            mime="text/markdown",
            use_container_width=True,
        )
    with d2:
        st.download_button(
            label="Download as Word (.docx)",
            data=markdown_to_docx_bytes(st.session_state.latest_output),
            file_name="application-pack.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

with st.expander("Session Learning Memory", expanded=False):
    if st.session_state.learning_notes:
        for i, note in enumerate(st.session_state.learning_notes[-20:], start=1):
            st.markdown(f"{i}. {note}")
    else:
        st.caption("No learning notes yet. Add rewrite feedback to build memory.")
